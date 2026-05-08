#!/usr/bin/env python3
"""
Generate a supplementary table in Word format, extracting the top 50 features
from both the baseline and slope elastic-net models.
"""

from pathlib import Path
import docx

# Import the logic and paths used in 'plot_feature_importance.py'
from plot_feature_importance import (
    load_and_prepare,
    FEATURE_IMPORTANCE_T1_CSV,
    FEATURE_IMPORTANCE_SLOPE_CSV
)

OUTPUT_DOCX = Path("/home/rachel/Desktop/superagers/elastic_net/figures_tables/supplementary_table.docx")

def process_top_50(df, is_slope_model):
    """Extract top 50 features by importance mean and tag annual change features.

    Args:
        df (pd.DataFrame): The feature importance DataFrame.
        is_slope_model (bool): Whether this is the slope model (for tagging).

    Returns:
        pd.DataFrame: Top 50 features sorted by descending importance.
    """
    top50 = df.sort_values("perm_importance_mean", ascending=False).head(50).copy()

    # Append (slope) label for longitudinal features
    if is_slope_model:
        top50.loc[top50["feature_type"] == "slope", "label"] = top50["label"] + " (annual change)"

    return top50


def main():
    # Load and prep data using the imported function
    df_baseline = load_and_prepare(FEATURE_IMPORTANCE_T1_CSV, is_slope_model=False)
    df_slope = load_and_prepare(FEATURE_IMPORTANCE_SLOPE_CSV, is_slope_model=True)

    top50_baseline = process_top_50(df_baseline, is_slope_model=False)
    top50_slope = process_top_50(df_slope, is_slope_model=True)

    doc = docx.Document()
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    headers = ["Feature", "Importance Mean ΔAUC"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    def add_model_section(sub_header_text, df):
        row = table.add_row()
        merged_cell = row.cells[0].merge(row.cells[1])
        merged_cell.text = sub_header_text
        merged_cell.paragraphs[0].runs[0].bold = True

        for _, r in df.iterrows():
            data_row = table.add_row()
            data_row.cells[0].text = f"    {r['label']}"
            data_row.cells[1].text = f"{r['perm_importance_mean']:.4f}"

    # Fill the document
    add_model_section("Baseline structure-function coupling model", top50_baseline)
    add_model_section("Baseline + annual change structure-function coupling model", top50_slope)

    doc.save(OUTPUT_DOCX)
    print(f"Saved supplementary table to: {OUTPUT_DOCX}")

if __name__ == "__main__":
    main()